import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io


def create_report(template_path, data, chart_data=None):
    st.write("Iniciando la creacion del informe")

    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f'{{{{{key}}}}}' in paragraph.text: # verificamos que la key exista en el documento
                st.write(f"Reemplazando {key} con {value} en el informe.")
            paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))
    
    if chart_data is not None:
        st.write("Generando grafico...")
        plt.figure(figsize=(6,4))
        plt.bar(chart_data['labels'], chart_data['values'])
        plt.title(chart_data['title'])
        plt.xlabel(chart_data.get('xlabel', ''))
        plt.ylabel(chart_data.get('ylabel', ''))
        
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        st.write("Insertando grafico en el marcador del documento...")
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if '[Aquí se insertara el grafico]' in run.text:
                    run.text = run.text.replace('[Aquí se insertara el grafico]', '')
                    run.add_picture(img_buffer, width=Inches(6))


    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    st.write("Informe creado con exito.")
    return output

def main():
    st.title("Generador de Informes desde Plantillas")
    template_file = st.file_uploader("Cargar plantilla Word", type="docx")
    data_file = st.file_uploader("Cargar datos", type=["xlsx", "csv"])

    if template_file and data_file:
        st.success("Archivos cargados correctamente")
        df = pd.read_csv(data_file) if data_file.name.endswith('.csv') else pd.read_excel(data_file)
        st.subheader("Datos cargados")
        st.dataframe(df)

        row_index = st.selectbox('Seleccinar fila para el informe', options=range(len(df)))
        selected_data = df.iloc[row_index].to_dict()
        
        generate_chart = st.checkbox("Generar grafico")
        chart_data = None

        if generate_chart:
            chart_title = st.text_input("Titulo del grafico", "Grafico de Datos")
            x_column = st.selectbox("Columna para eje X", options=df.columns)
            y_column = st.selectbox("Columna para eje Y", options=df.columns)
            print(df.columns)
            chart_data = {
                'title': chart_title,
                'labels': df[x_column].tolist(),
                'values': df[y_column].tolist(),
                'xlabel': x_column,
                'ylabel': y_column
            }

            st.write("Datos del grafico:", chart_data)

        if st.button("Generar Informe"):
            output = create_report(template_file, selected_data, chart_data)
            st.download_button("Descargar Informe", output, "informe_generado.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        
main()